import datetime
from lxml import html
import urllib.request
from bs4 import BeautifulSoup
# from bs4 import SoupStrainer # to use this, add 'parse_only="something" to BeautifulSoup()
# and it will only grab those parts (e.g. parse_only="b" to show bold stuff)
from docx import Document
from os import path

# Settings
#import "config.py"

# These are now in config.py
baseurl = 'http://le.utah.gov/'
#filepath = 'E:\\!Morgan\\CC\\Legislative Sessions\\2017 Bill Outputs'
filepath = 'C:\\Temp\\2018 Bill Outputs'
#if not path.exists(filepath):
#    filepath = input('Uhm... \'{}\' doesn\'t exist, where would you like the file(s) saved? '.format(filepath))

print('Please enter all bill IDs you want formatted. If you have more than one, separate them with commas.')
str_bills = input('Bill(s): ')

#bills = ['HB0241', 'HB0291', 'HB0249']  # HB0245, hb0291,HB0249
bills = str_bills.replace(' ', '').upper().split(',')

for bill_id in bills:
    if len(bill_id) < 6:
        bill_id = '{}{:0>4}'.format(bill_id[:2],bill_id[2:])
    print('Processing bill %s' % (bill_id))
    document = Document()

    thisyear = datetime.datetime.now().year
    #fullurl = '%s~%s/bills/static/%s.html' % (baseurl, thisyear, bill_id) #might want to try pulling from source page instead of "printer friendly" to automatically get the latest? (find a way to get older versions too?)
    fullurl = '%s~%s/bills/%sbillint/%s.htm' % (baseurl, thisyear, bill_id[0].lower(), bill_id) #the original pulled from "printer friendly" version
    print(' > downloading %s' % (fullurl))
    with urllib.request.urlopen(fullurl) as response:
        soup = BeautifulSoup(response.read(), 'lxml')
        body = soup.body.div.b.find_all_next(string=True)
        bill = soup.title.string.replace('Utah Legislature ', '')
        header = body[0]
        for item in body:
            if item.startswith('Chief Sponsor:'):
                sponsor = item.replace('Chief Sponsor:  ', '')
        document.add_heading(bill, level=0)
        print(' > building header')
        document.add_heading('1. %s %s (%s)' % (bill, header, sponsor), level=3)

        print(' > extracting description and highlights')
        section = 0
        gen_desc = ''
        this_bill = []
        for line in body:
            if section == 0:
                if line == 'General Description:':
                    section = 1
            elif section == 1:
                if line[-10:] == 'This bill:':
                    section = 2
                elif line != 'Highlighted Provisions:':
                    gen_desc = gen_desc + ' ' + line[2:].lstrip('\xa0').rstrip('\xa0')
            elif section == 2:
                if line == 'Money Appropriated in this Bill:':
                    section = 3
                else:
                    this_bill.append(line[2:].lstrip('\xa0').rstrip('\xa0'))
            elif section == 3:
                break

    print(' > formatting description')
    document.add_paragraph('General Description:', style='List Bullet')
    document.add_paragraph(gen_desc[1:], style='List Bullet 2')

    print(' > formatting highlights')
    document.add_paragraph('This bill:', style='List Bullet')
    bullets = []
    for item in this_bill:
        if any(c in item for c in ('▸','•')):
            bullets.append(item[0] + item[1:].lstrip('\xa0'))
        else:
            bullets[-1] = bullets[-1] + ' ' + item
    for item in bullets:
        if item[0] == '▸':
            document.add_paragraph(item[1:],style='List Bullet 2')
        elif item[0] == '•':
            document.add_paragraph(item[1:],style='List Bullet 3')
        else:
            document.add_paragraph(item)

    print(' > writing Word document')
    try:
        document.save('%s\\%s.docx' % (filepath, bill_id))
    except FileNotFoundError:
        print('\'{}\' does not exist, could not save {}.docx'.format(filepath,bill_id))
input('Press <Enter> to close this window...')